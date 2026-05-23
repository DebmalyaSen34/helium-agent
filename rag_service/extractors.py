from __future__ import annotations

import csv
import hashlib
import io
import json
import re
from pathlib import Path

from rag_service.models import ExtractedBlock, ExtractedDocument, RagError, SourceLocation


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".log",
}


def file_hash(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def extract_document(path: Path) -> ExtractedDocument:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return extract_text_document(path)
    if suffix == ".csv":
        return extract_csv_document(path)
    if suffix == ".pdf":
        return extract_pdf_document(path)
    if suffix == ".docx":
        return extract_docx_document(path)
    if suffix == ".xlsx":
        return extract_xlsx_document(path)
    raise RagError("unsupported_file", f"Unsupported attachment type: {suffix or '(none)'}.")


def extract_text_document(path: Path) -> ExtractedDocument:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as exc:
        raise RagError("decode_error", "Attachment is not valid UTF-8 text.") from exc

    lines = text.splitlines()
    blocks: list[ExtractedBlock] = []
    outline: list[str] = []
    suffix = path.suffix.lower()

    if suffix == ".md":
        blocks, outline = _markdown_blocks(lines)
    elif suffix == ".py":
        blocks, outline = _python_blocks(lines)
    else:
        blocks = [
            ExtractedBlock(
                text=text,
                location=SourceLocation(kind="lines", label="full file", start_line=1, end_line=max(1, len(lines))),
                title="full file",
            )
        ]

    return ExtractedDocument(
        file_path=str(path),
        file_name=path.name,
        file_hash=file_hash(path),
        byte_size=path.stat().st_size,
        kind=suffix.lstrip(".") or "text",
        text=text,
        blocks=tuple(block for block in blocks if block.text.strip()),
        outline=tuple(outline),
    )


def _markdown_blocks(lines: list[str]) -> tuple[list[ExtractedBlock], list[str]]:
    heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
    starts: list[tuple[int, str]] = []
    outline: list[str] = []
    for index, line in enumerate(lines, start=1):
        match = heading_re.match(line)
        if match:
            title = match.group(2).strip()
            starts.append((index, title))
            outline.append(f"{match.group(1)} {title}")

    if not starts:
        text = "\n".join(lines)
        return [
            ExtractedBlock(
                text=text,
                location=SourceLocation(kind="lines", label="full file", start_line=1, end_line=max(1, len(lines))),
                title="full file",
            )
        ], outline

    blocks: list[ExtractedBlock] = []
    for idx, (start, title) in enumerate(starts):
        end = starts[idx + 1][0] - 1 if idx + 1 < len(starts) else len(lines)
        text = "\n".join(lines[start - 1:end])
        blocks.append(
            ExtractedBlock(
                text=text,
                location=SourceLocation(kind="lines", label=title, start_line=start, end_line=end),
                title=title,
            )
        )
    return blocks, outline


def _python_blocks(lines: list[str]) -> tuple[list[ExtractedBlock], list[str]]:
    try:
        import ast

        tree = ast.parse("\n".join(lines))
    except SyntaxError:
        text = "\n".join(lines)
        return [
            ExtractedBlock(
                text=text,
                location=SourceLocation(kind="lines", label="full file", start_line=1, end_line=max(1, len(lines))),
                title="full file",
            )
        ], ()

    blocks: list[ExtractedBlock] = []
    outline: list[str] = []
    imports: list[str] = []
    import_start: int | None = None
    import_end: int | None = None

    for node in tree.body:
        if isinstance(node, (ast.Import, ast.ImportFrom)):
            imports.append("\n".join(lines[node.lineno - 1:getattr(node, "end_lineno", node.lineno)]))
            import_start = node.lineno if import_start is None else min(import_start, node.lineno)
            import_end = max(import_end or node.lineno, getattr(node, "end_lineno", node.lineno))
            continue
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            start = node.lineno
            end = getattr(node, "end_lineno", node.lineno)
            kind = "class" if isinstance(node, ast.ClassDef) else "function"
            title = f"{kind} {node.name}"
            outline.append(title)
            blocks.append(
                ExtractedBlock(
                    text="\n".join(lines[start - 1:end]),
                    location=SourceLocation(kind="lines", label=title, start_line=start, end_line=end),
                    title=title,
                )
            )

    if imports and import_start is not None and import_end is not None:
        blocks.insert(
            0,
            ExtractedBlock(
                text="\n".join(imports),
                location=SourceLocation(kind="lines", label="imports", start_line=import_start, end_line=import_end),
                title="imports",
            ),
        )
        outline.insert(0, "imports")

    if not blocks:
        text = "\n".join(lines)
        blocks.append(
            ExtractedBlock(
                text=text,
                location=SourceLocation(kind="lines", label="full file", start_line=1, end_line=max(1, len(lines))),
                title="full file",
            )
        )
    return blocks, outline


def extract_csv_document(path: Path) -> ExtractedDocument:
    text = path.read_text(encoding="utf-8-sig")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        raise RagError("empty_file", "CSV attachment is empty.")

    headers = rows[0]
    sample = rows[1:11]
    summary = {
        "sheet": "CSV",
        "headers": headers,
        "row_count": max(0, len(rows) - 1),
        "column_count": len(headers),
        "sample_rows": sample,
    }
    rendered = json.dumps(summary, indent=2, ensure_ascii=False)
    block = ExtractedBlock(
        text=rendered,
        location=SourceLocation(kind="sheet", label="CSV", sheet="CSV", start_row=1, end_row=min(len(rows), 11)),
        title="CSV summary",
    )
    return ExtractedDocument(
        file_path=str(path),
        file_name=path.name,
        file_hash=file_hash(path),
        byte_size=path.stat().st_size,
        kind="csv",
        text=rendered,
        blocks=(block,),
        outline=(f"CSV: {len(headers)} columns, {max(0, len(rows) - 1)} data rows",),
    )


def extract_pdf_document(path: Path) -> ExtractedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RagError("missing_dependency", "PDF extraction requires pypdf. Install requirements-rag.txt.") from exc

    reader = PdfReader(str(path))
    blocks: list[ExtractedBlock] = []
    all_text: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            all_text.append(page_text)
            blocks.append(
                ExtractedBlock(
                    text=page_text,
                    location=SourceLocation(kind="page", label=f"page {index}", page=index),
                    title=f"page {index}",
                )
            )

    text = "\n\n".join(all_text).strip()
    if len(text) < 80:
        raise RagError("ocr_not_supported", "PDF has little or no extractable text. OCR is not supported in v1.")

    return ExtractedDocument(
        file_path=str(path),
        file_name=path.name,
        file_hash=file_hash(path),
        byte_size=path.stat().st_size,
        kind="pdf",
        text=text,
        blocks=tuple(blocks),
        outline=tuple(block.title for block in blocks),
    )


def extract_docx_document(path: Path) -> ExtractedDocument:
    try:
        from docx import Document
    except ImportError as exc:
        raise RagError("missing_dependency", "DOCX extraction requires python-docx. Install requirements-rag.txt.") from exc

    document = Document(str(path))
    blocks: list[ExtractedBlock] = []
    current: list[str] = []
    current_title = "document"
    start_line = 1
    line_no = 0
    outline: list[str] = []

    def flush() -> None:
        nonlocal current, start_line, current_title
        if current:
            blocks.append(
                ExtractedBlock(
                    text="\n".join(current),
                    location=SourceLocation(kind="lines", label=current_title, start_line=start_line, end_line=line_no),
                    title=current_title,
                )
            )
            current = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if not value:
            continue
        line_no += 1
        style = paragraph.style.name.lower() if paragraph.style else ""
        if "heading" in style:
            flush()
            current_title = value
            outline.append(value)
            start_line = line_no
        current.append(value)

    flush()
    text = "\n\n".join(block.text for block in blocks)
    if not text.strip():
        raise RagError("empty_file", "DOCX attachment did not contain extractable text.")

    return ExtractedDocument(
        file_path=str(path),
        file_name=path.name,
        file_hash=file_hash(path),
        byte_size=path.stat().st_size,
        kind="docx",
        text=text,
        blocks=tuple(blocks),
        outline=tuple(outline),
    )


def extract_xlsx_document(path: Path) -> ExtractedDocument:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RagError("missing_dependency", "XLSX extraction requires openpyxl. Install requirements-rag.txt.") from exc

    workbook = load_workbook(str(path), read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    outline: list[str] = []
    rendered_sheets: list[str] = []

    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(value) if value is not None else "" for value in rows[0]]
        sample_rows = [[value for value in row] for row in rows[1:11]]
        summary = {
            "sheet": sheet.title,
            "headers": headers,
            "row_count": max(0, len(rows) - 1),
            "column_count": len(headers),
            "sample_rows": sample_rows,
        }
        rendered = json.dumps(summary, indent=2, ensure_ascii=False, default=str)
        rendered_sheets.append(rendered)
        outline.append(f'{sheet.title}: {len(headers)} columns, {max(0, len(rows) - 1)} data rows')
        blocks.append(
            ExtractedBlock(
                text=rendered,
                location=SourceLocation(
                    kind="sheet",
                    label=sheet.title,
                    sheet=sheet.title,
                    start_row=1,
                    end_row=min(len(rows), 11),
                ),
                title=f"{sheet.title} summary",
            )
        )

    if not blocks:
        raise RagError("empty_file", "XLSX attachment did not contain extractable sheets.")

    return ExtractedDocument(
        file_path=str(path),
        file_name=path.name,
        file_hash=file_hash(path),
        byte_size=path.stat().st_size,
        kind="xlsx",
        text="\n\n".join(rendered_sheets),
        blocks=tuple(blocks),
        outline=tuple(outline),
    )
